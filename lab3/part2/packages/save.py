from docx import Document
from docx.shared import Inches
import os.path


class SaveCha:
    def __init__(self):
        self.path = os.path.dirname(__file__)
        self.fullaPath = os.path.join(self.path, 'Simple.docx')
        if os.path.exists(self.fullaPath):
            self.document = Document('Simple.docx')
        else:
            self.document = Document()
    def SaveChan(self, electron, proton, neitron):

        self.document.add_paragraph(
            "Электрон:  \t" + electron, style='List Bullet')
        self.document.add_paragraph(
            "Протон:  \t" + proton, style='List Bullet')
        self.document.add_paragraph(
            "Нейтрон:  \t" + neitron, style='List Bullet')

        self.document.add_paragraph("\n")

        self.document.save('Simple.docx')
